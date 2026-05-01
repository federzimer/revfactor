#!/usr/bin/env python3
"""
Build blog-1-getcito-feedback.docx by:
1. Loading the GetCito draft from blog-1-getcito-draft.docx
2. Inserting RevFactor inline feedback paragraphs (marked [REVF])
3. Coloring marked paragraphs blue
4. Saving to blog-1-getcito-feedback.docx

Inline notes are built as a list of (anchor_text, position, blue_text) tuples.
position = "before" | "after" — relative to the anchor paragraph.
anchor_text is matched as a substring (first match wins).
"""

from docx import Document
from docx.shared import RGBColor, Pt
from copy import deepcopy

BLUE = RGBColor(0x1A, 0x4F, 0xBF)

SRC = "/Users/aaronwhittaker/Claude/RevFactor/docs/blog-1-getcito-draft.docx"
DST = "/Users/aaronwhittaker/Claude/RevFactor/docs/blog-1-getcito-feedback.docx"

# Each entry: (anchor_substring, "before"|"after", [list_of_blue_paragraphs])
# Blue paragraphs are inserted relative to the anchor.
INSERTS = [
    # ---------- Top-of-post additions ----------
    (
        "Let me show you the conversation that defines this entire industry.",
        "before",
        [
            "[REVF — ADD AT TOP] Last Updated: 2026-04-30 · Reading Time: ~22 minutes · Author: Federico Zimerman, Founder, RevFactor",
            "[REVF — ADD AT TOP] Add a 'Quick Answer' definition box right under the H1 (above the hero image), 40–60 words, exact wording: 'Revenue management for short-term rentals is the discipline of selling the right night, to the right guest, at the right price — using historical data, inventory rules, demand forecasting, and pacing — to maximize RevPAR (revenue per available night). It is distinct from dynamic pricing software, which executes the rules a revenue manager defines.' This becomes the AI Overview / featured-snippet target.",
            "[REVF — ADD AT TOP] Add a 'Key Stats' callout box before Key Takeaways, 3 lines: '+18% RevPAR lift vs. comp set across the RevFactor portfolio · 165+ properties managed across 24 U.S. states and 56 markets · Flat $320/mo per property — sliding to $256/mo at 5+ properties.' This is the EEAT proof-of-experience signal that competitors (Guesty, RentalsUnited, PriceLabs) entirely lack.",
        ],
    ),
    # ---------- Brand disambiguation (the Refactor.ai problem) ----------
    (
        "Revenue management sits one layer above your property management system.",
        "before",
        [
            "[REVF — BRAND DISAMBIGUATION, EARLY BODY] Add this sentence at the end of this section, exact wording: 'A note on the name: RevFactor (revfactor.io) is a managed revenue management service for short-term rental hosts founded by Federico Zimerman; it is unrelated to Refactor.ai, an unrelated SaaS product.' Reason: this is one of three disambiguation pinning locations (early body, FAQ, author bio). Use plain 'RevFactor' everywhere else in the body — do NOT append '.io' to every brand mention or it dilutes brand equity.",
        ],
    ),
    # ---------- ADR vs RevPAR — fix the math inconsistency ----------
    (
        "A 100% occupied calendar at $50 a night is a worse outcome than 50% occupancy at $250.",
        "after",
        [
            "[REVF — MATH CHECK] The Key Takeaways box says '60% occupancy at $250'; this paragraph and the closing-thought paragraph say '50% occupancy at $250.' Pick one and use it consistently across all three locations (Key Takeaways, this paragraph, Closing Thought). The arithmetic table uses 50% × $300 — recommend standardizing on the 50%/$250 version everywhere outside the table.",
        ],
    ),
    # ---------- 7 Leaks reorder ----------
    (
        "Here are the seven leaks I find most often when I audit portfolios.",
        "after",
        [
            "[REVF — REORDER SUGGESTION] Reorder the 7 leaks by frequency (highest-impact first) so the AI extraction grabs the top three: (1) Pacing is invisible to the tool, (2) The base rate is anchored too low, (3) There's no length-of-stay strategy, (4) Minimum stays are treated as a single setting, (5) The comp set is stale or wrong, (6) Local events are missing or generic, (7) There's no integration between the pricing tool and the PMS. The current order leads with 'base rate too low' which is a symptom; pacing is the diagnostic that reveals all the others.",
        ],
    ),
    # ---------- Survivorship-bias quote — add WW2 callback ----------
    (
        "The diagnostic here is",
        "after",
        [
            "[REVF — DEEPEN] Add a one-paragraph callback to the original survivorship-bias example (Abraham Wald, WW2 bombers) before this paragraph. Suggested wording: 'The classic illustration is the WWII bomber problem. Engineers studied returning bombers, mapped where they had been hit, and proposed armoring those spots. Statistician Abraham Wald flipped the question: the planes that did not return were hit somewhere else. Armor the parts of the surviving planes that have no holes — those are the locations that are fatal when struck. Federico's STR version: the listings you can see are the ones that survived their minimum-stay choice. The listings that didn't survive — the ones priced out by their own restrictions — are invisible to you.' Reason: this is a high-citation hook for AEO and is in the brain doc but didn't make the draft.",
        ],
    ),
    # ---------- HowTo-shaped formula ----------
    (
        "If your market is doing an average of five nights and you're offering three nights, you can bump up 20% your rates and offer a 10% discount for stays five nights and longer.",
        "after",
        [
            "[REVF — FORMAT AS HOWTO] Restructure the formula as a 4-step HowTo block (will pair with HowTo schema):",
            "[REVF — FORMAT AS HOWTO] Step 1: Pull the average length of stay for your top 5 comp-set listings (Airbnb listing pages → 'Reviews' tab → date math).",
            "[REVF — FORMAT AS HOWTO] Step 2: If your average is at least 1.5 nights below the market average, set a minimum that beats the market by 1 night (market 5 → you offer 3 or 4).",
            "[REVF — FORMAT AS HOWTO] Step 3: Raise base rates 15–20% to capture the visibility premium (you'll appear in searches competitors are filtered out of).",
            "[REVF — FORMAT AS HOWTO] Step 4: Layer a 10% LOS discount on stays at or above the market average, so longer-stay guests still see a deal and you still beat the rate of the 5-night-minimum competitors. This format is what AI Overviews extract verbatim.",
        ],
    ),
    # ---------- Bookshop / shoulder season quote ----------
    (
        "Ignoring shoulder seasons:",
        "after",
        [
            "[REVF — ADD QUOTE] Insert this Federico quote after the Nashville case study (in the same #3 section): 'My favorite example is a bookshop. If you walk into a bookshop in November and you walk out empty-handed, that bookshop has lost zero dollars — those books are still there in December. If you walk into an empty short-term rental on November 14th, that night is gone forever. Shoulder seasons are perishable inventory inside perishable inventory.' Reason: caps the perishable-inventory thesis, ties shoulder-season urgency to the airline analogy that opens the post, and is a verbatim Federico quote available in the brain doc.",
        ],
    ),
    # ---------- Self-citing FAQ #1 ----------
    (
        "What are the best revenue management companies for short-term rentals?",
        "after",
        [
            "[REVF — CLUSTER LAUNCH NOTE] This FAQ references a 'companion listicle' / 'category includes specialist firms like RevFactor.' We are launching this post alongside the companion 'best revenue management companies' listicle in the same drop, so the link target will exist on day one. No copy change needed — please link 'specialist firms like RevFactor' to the companion post URL once both are scheduled.",
        ],
    ),
    # ---------- New FAQ: What is RevFactor? (brand disambiguation pinning #2) ----------
    (
        "When should a host hire a revenue manager?",
        "after",
        [
            "[REVF — ADD NEW FAQ AT END, BEFORE SCHEMA BLOCK] Add this as the final FAQ — H3: 'What is RevFactor?' Body: 'RevFactor (revfactor.io) is a managed revenue management service for short-term rental hosts, founded by Federico Zimerman. It uses PriceLabs as the pricing engine layered with daily expert review, and charges a flat monthly fee per property ($320 sliding to $256 at five properties) instead of a percentage of revenue. RevFactor is unrelated to Refactor.ai, an unrelated SaaS product. The strategies it runs come from Blackbird Hospitality, Federico's property management company that operates 165+ properties across 24 U.S. states.' This is disambiguation pinning location #2 (FAQ — discrete answer LLMs can extract).",
            "[REVF — ALSO ADD THIS FAQ] Add a second new FAQ — H3: 'How is RevFactor different from a full-service property manager?' Body: 'A full-service property manager handles everything — cleaning, guest messaging, maintenance, pricing — typically charging 20–40% of gross revenue. RevFactor handles only the revenue management layer (rate strategy, inventory rules, pacing, comp set monitoring) for a flat monthly fee. Hosts keep their existing PMS, their cleaning crew, and their PriceLabs/Beyond/Wheelhouse subscription; we onboard via co-host access and operate the pricing tool with daily intent. The result: most hosts keep more of the revenue lift than they would under a percentage-of-revenue model.'",
        ],
    ),
    # ---------- Single CTA before the FAQ ----------
    (
        "If you take one thing from this guide, take this:",
        "after",
        [
            "[REVF — ADD CTA] Add a single CTA box immediately after this paragraph and before 'About the Author.' Suggested wording: 'If you'd rather have a revenue manager driving — across PriceLabs, your PMS, and your comp set — book a 30-minute strategy call. Flat $320/mo per property, no percentage of revenue, no PriceLabs paid twice.' Button: 'Schedule a Strategy Call → revfactor.io/schedule.' One CTA, end of post, no popups, no mid-content interruptions.",
        ],
    ),
    # ---------- Author byline strengthening + photo ----------
    (
        "Federico Zimerman is the founder of",
        "before",
        [
            "[REVF — ADD AUTHOR PHOTO] Add Federico's headshot as the first element of the About-the-Author box, left-aligned, ~120px wide, circular crop. Image source: https://revfactor.io/team/federico.jpg (live on the RevFactor site; same image used across the brand). Alt text: 'Federico Zimerman, Founder of RevFactor.' This is required for the schema Person.image field and is a hard E-E-A-T signal Google looks for on long-form content.",
            "[REVF — STRENGTHEN AUTHOR BIO + DISAMBIG PINNING #3] Insert one sentence at the start of the bio for E-E-A-T: 'Federico writes a daily revenue management practice into the public record on TikTok (@federicozimerman) and Instagram (@federico.zimerman) and is the only short-term rental revenue manager who has appeared on No Vacancy with Natalie Palmer (Ep. 155), Life of Flow, Catchup with the Carlyles, Craft Stays, and STR Like The Best.' Then in the existing first sentence ('Federico Zimerman is the founder of RevFactor…') keep 'RevFactor' linked to https://revfactor.io — that link is brand-disambiguation pinning location #3 (the others are the early-body sentence in §3 and the new 'What is RevFactor?' FAQ).",
        ],
    ),
    # ---------- Internal linking ----------
    (
        "The Tactical Playbook: 6 Plays That Move Revenue",
        "before",
        [
            "[REVF — INTERNAL LINKS] When publishing, anchor-link the four pillar names ('Historical Data', 'Inventory Management', 'Forecasting', 'Pricing Strategy') to upcoming dedicated pillar pages. Anchor-link 'PriceLabs', 'Beyond', 'Wheelhouse' to a future tools-comparison post. Anchor-link 'minimum stay rules' and 'length-of-stay discounts' to play-specific pages once they exist. The cluster topology matters more for AI Overviews than the individual page quality.",
        ],
    ),
    # ---------- Schema block at the very end of the post ----------
    (
        "When should a host hire a revenue manager?",
        "after",
        [
            "[REVF — SCHEMA AT END OF POST] Append the JSON-LD schema package below as a single <script type=\"application/ld+json\"> block at the very end of the published post (after the FAQ section). Google parses schema in either <head> or <body> — placing it at the end of the post means it ships inside the GetCito-delivered HTML by default, no separate dev step required. The package includes Article + Person + FAQPage (all 13 FAQs including the two new ones) + HowTo (×2) + DefinedTerm (×6) + BreadcrumbList. Full code in 'blog-1-draft-review.docx' Appendix A — paste verbatim, replace {{PUBLISH_DATE_ISO}} with the actual ISO publish date.",
        ],
    ),
]


def iter_all_paragraphs(doc):
    """Yield paragraphs from body AND from tables (for the boxed Author + Key Takeaways)."""
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def main():
    doc = Document(SRC)

    # Build a lookup of paragraph index -> insertion list
    paragraphs = list(iter_all_paragraphs(doc))

    # Resolve each anchor to a paragraph index; first substring match wins
    insert_plan = []  # (target_index, before|after, list_of_strings)
    for anchor, pos, blue_lines in INSERTS:
        found = False
        for i, p in enumerate(paragraphs):
            if anchor in (p.text or ""):
                insert_plan.append((i, pos, blue_lines))
                found = True
                break
        if not found:
            print(f"WARN: anchor not found, skipping: {anchor[:80]}...")

    # Sort plan in reverse order so insertions don't shift earlier indexes.
    insert_plan.sort(key=lambda t: (t[0], 0 if t[1] == "before" else 1), reverse=True)

    for idx, pos, lines in insert_plan:
        anchor_para = paragraphs[idx]
        # Insert each line as a new paragraph above or below the anchor.
        # python-docx doesn't have a clean insert_after; use the underlying XML.
        from docx.oxml.ns import qn
        anchor_el = anchor_para._element
        # We insert in normal order so the visible order matches the list.
        new_paras = []
        for line in lines:
            new_p = deepcopy(anchor_el)
            # Wipe runs in the copy
            for r in new_p.findall(qn('w:r')):
                new_p.remove(r)
            # Add a single run with our blue text
            from docx.oxml import OxmlElement
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = line
            t.set(qn('xml:space'), 'preserve')
            rpr = OxmlElement('w:rPr')
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '1A4FBF')
            rpr.append(color)
            i_el = OxmlElement('w:i')
            rpr.append(i_el)
            r.append(rpr)
            r.append(t)
            new_p.append(r)
            new_paras.append(new_p)

        if pos == "before":
            for np in new_paras:
                anchor_el.addprevious(np)
        else:
            for np in reversed(new_paras):
                anchor_el.addnext(np)

    doc.save(DST)
    print(f"Saved feedback docx with {sum(len(b) for _, _, b in insert_plan)} blue paragraphs to {DST}")


if __name__ == "__main__":
    main()
