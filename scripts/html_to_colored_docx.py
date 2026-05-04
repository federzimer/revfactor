#!/usr/bin/env python3
"""
Convert blog-N-outline-feedback.html → .docx with color-coded runs.

Color rules (matches HTML CSS):
  .original / outside-callouts → black (brief content)
  .blue / .blue.banner         → blue   #1A4FBF (Claude/Aaron feedback)
  .red                         → red    #8B1A1A (critical warnings)

Pass HTML path as argv[1]. Writes .docx alongside.
"""
import sys
from pathlib import Path
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0x1A, 0x4F, 0xBF)
BLUE_DARK = RGBColor(0x0A, 0x2E, 0x80)
RED = RGBColor(0x8B, 0x1A, 0x1A)
RED_DARK = RGBColor(0x5A, 0x00, 0x00)
BLACK = RGBColor(0x22, 0x22, 0x22)
GREEN = RGBColor(0x0A, 0x7A, 0x3C)
GREEN_DARK = RGBColor(0x0A, 0x4F, 0x25)
GRAY = RGBColor(0x88, 0x88, 0x88)

BLUE_BG_HEX = "F0F5FF"
RED_BG_HEX = "FFF0F0"
GREEN_BG_HEX = "F0FFF5"
CODE_BG_HEX = "F6F8FA"

CTX_TO_COLOR = {"blue": BLUE, "red": RED, "green": GREEN, "black": BLACK}
CTX_TO_BG = {"blue": BLUE_BG_HEX, "red": RED_BG_HEX, "green": GREEN_BG_HEX}
CTX_TO_BORDER = {"blue": "1A4FBF", "red": "8B1A1A", "green": "0A7A3C"}
CTX_TO_LABEL_COLOR = {"blue": BLUE_DARK, "red": RED_DARK, "green": GREEN_DARK}


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def shade_paragraph(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def add_left_border(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def context_color(ctx):
    return CTX_TO_COLOR.get(ctx, BLACK)


def add_runs(paragraph, node, ctx, bold=False, italic=False):
    """Recursively walk inline content of `node`, appending runs to `paragraph`."""
    color = context_color(ctx)
    if isinstance(node, NavigableString):
        text = str(node)
        if not text:
            return
        run = paragraph.add_run(text)
        run.font.color.rgb = color
        run.font.size = Pt(11)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        return
    if not isinstance(node, Tag):
        return
    name = node.name
    if name in ("strong", "b"):
        for child in node.children:
            add_runs(paragraph, child, ctx, bold=True, italic=italic)
    elif name in ("em", "i"):
        for child in node.children:
            add_runs(paragraph, child, ctx, bold=bold, italic=True)
    elif name == "br":
        run = paragraph.add_run()
        run.add_break()
    elif name == "code":
        run = paragraph.add_run(node.get_text())
        run.font.color.rgb = color
        run.font.name = "Menlo"
        run.font.size = Pt(10)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
    elif name == "a":
        text = node.get_text()
        href = node.get("href", "")
        run = paragraph.add_run(text)
        run.font.color.rgb = color
        run.font.size = Pt(11)
        run.font.underline = True
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if href:
            run.add_comment if False else None
    elif name == "span":
        classes = node.get("class", [])
        # Track how many runs we add so we can re-style only those
        before_count = len(paragraph.runs)
        for child in node.children:
            add_runs(paragraph, child, ctx, bold=bold, italic=italic)
        added_runs = paragraph.runs[before_count:]
        if "strike" in classes:
            for run in added_runs:
                run.font.strike = True
                run.font.color.rgb = GRAY
        elif "resolved" in classes or "green-text" in classes:
            for run in added_runs:
                run.font.color.rgb = GREEN
                run.bold = True
    else:
        for child in node.children:
            add_runs(paragraph, child, ctx, bold=bold, italic=italic)


def render_paragraph(doc, p_tag, ctx, style=None, shade=False, border=False):
    paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if shade and ctx == "blue":
        shade_paragraph(paragraph, BLUE_BG_HEX)
    if shade and ctx == "red":
        shade_paragraph(paragraph, RED_BG_HEX)
    if border and ctx == "blue":
        add_left_border(paragraph, "1A4FBF")
    if border and ctx == "red":
        add_left_border(paragraph, "8B1A1A")
    for child in p_tag.children:
        add_runs(paragraph, child, ctx)
    return paragraph


def render_list(doc, list_tag, ctx, ordered=False):
    style = "List Number" if ordered else "List Bullet"
    for li in list_tag.find_all("li", recursive=False):
        p = doc.add_paragraph(style=style)
        for child in li.children:
            # Skip nested lists from being inlined; render them after
            if isinstance(child, Tag) and child.name in ("ul", "ol"):
                continue
            add_runs(p, child, ctx)
        # Recurse for nested lists
        for nested in li.find_all(["ul", "ol"], recursive=False):
            render_list(doc, nested, ctx, ordered=(nested.name == "ol"))


def render_table(doc, table_tag, ctx):
    rows = table_tag.find_all("tr", recursive=False)
    if not rows:
        # Try inside thead/tbody
        rows = table_tag.find_all("tr")
    if not rows:
        return
    cols = max(len(r.find_all(["td", "th"], recursive=False)) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid"
    for ri, tr in enumerate(rows):
        cells = tr.find_all(["td", "th"], recursive=False)
        is_header = any(c.name == "th" for c in cells)
        for ci, cell_tag in enumerate(cells):
            if ci >= cols:
                break
            cell = table.rows[ri].cells[ci]
            cell.text = ""  # clear default empty paragraph text
            p = cell.paragraphs[0]
            for child in cell_tag.children:
                add_runs(p, child, ctx, bold=is_header)
            if ctx == "blue":
                shade_cell(cell, "E0EAFF" if is_header else BLUE_BG_HEX)
            elif ctx == "red":
                shade_cell(cell, RED_BG_HEX)
            elif is_header:
                shade_cell(cell, "F4F4F4")
    doc.add_paragraph()  # spacer


def render_callout_block(doc, div, ctx):
    """Render a div with class .blue / .red / .green as a sequence of shaded paragraphs."""
    bg = CTX_TO_BG.get(ctx, BLUE_BG_HEX)
    border = CTX_TO_BORDER.get(ctx, "1A4FBF")
    label_color = CTX_TO_LABEL_COLOR.get(ctx, BLUE_DARK)
    for child in div.children:
        if isinstance(child, NavigableString):
            txt = str(child).strip()
            if not txt:
                continue
            p = doc.add_paragraph()
            shade_paragraph(p, bg)
            add_left_border(p, border)
            run = p.add_run(txt)
            run.font.color.rgb = context_color(ctx)
            run.font.size = Pt(11)
            continue
        if not isinstance(child, Tag):
            continue
        name = child.name
        classes = child.get("class", [])
        if name == "p":
            is_label = "label" in classes or "label-red" in classes
            p = doc.add_paragraph()
            shade_paragraph(p, bg)
            add_left_border(p, border)
            for inline in child.children:
                add_runs(p, inline, ctx, bold=is_label)
            if is_label:
                for run in p.runs:
                    run.text = run.text.upper()
                    run.font.size = Pt(9)
                    run.font.color.rgb = label_color
        elif name in ("ul", "ol"):
            render_list(doc, child, ctx, ordered=(name == "ol"))
        elif name == "table":
            render_table(doc, child, ctx)
        elif name == "blockquote":
            p = doc.add_paragraph()
            shade_paragraph(p, "FFFFFF")
            add_left_border(p, border)
            for inline in child.children:
                add_runs(p, inline, ctx, italic=True)
        elif name == "pre":
            # Code block — render as monospace, light gray bg, smaller font
            text = child.get_text()
            for line in text.split("\n"):
                p = doc.add_paragraph()
                shade_paragraph(p, CODE_BG_HEX)
                run = p.add_run(line if line.strip() else " ")
                run.font.name = "Menlo"
                run.font.size = Pt(9)
                run.font.color.rgb = BLACK
        elif name == "div":
            sub_classes = child.get("class", [])
            if "blue" in sub_classes:
                render_callout_block(doc, child, "blue")
            elif "red" in sub_classes:
                render_callout_block(doc, child, "red")
            elif "green" in sub_classes:
                render_callout_block(doc, child, "green")
            else:
                render_callout_block(doc, child, ctx)
        else:
            p = doc.add_paragraph()
            shade_paragraph(p, bg)
            for inline in child.children:
                add_runs(p, inline, ctx)


def convert(html_path: Path, out_path: Path):
    soup = BeautifulSoup(html_path.read_text(), "html.parser")
    body = soup.body
    doc = Document()

    # Wider page margins for readability
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # Set default Normal font
    style = doc.styles["Normal"]
    style.font.name = "Helvetica"
    style.font.size = Pt(11)

    # Color key legend — explicit so the reader doesn't have to ask what each color means
    legend = doc.add_paragraph()
    legend_run = legend.add_run("COLOR KEY: ")
    legend_run.bold = True
    legend_run.font.size = Pt(9)
    r1 = legend.add_run("Black = original brief  ·  ")
    r1.font.color.rgb = BLACK
    r1.font.size = Pt(9)
    r2 = legend.add_run("Green = feedback  ·  ")
    r2.font.color.rgb = GREEN
    r2.font.size = Pt(9)
    r2.bold = True
    r3 = legend.add_run("Red = critical / blocker")
    r3.font.color.rgb = RED
    r3.font.size = Pt(9)
    r3.bold = True

    for child in body.children:
        if isinstance(child, NavigableString):
            txt = str(child).strip()
            if txt:
                doc.add_paragraph(txt)
            continue
        if not isinstance(child, Tag):
            continue
        name = child.name
        classes = child.get("class", [])
        if name == "h1":
            p = doc.add_paragraph()
            run = p.add_run(child.get_text())
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = BLACK
        elif name == "h2":
            p = doc.add_paragraph()
            run = p.add_run(child.get_text())
            run.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = BLACK
            # bottom rule
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "999999")
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif name == "h3":
            p = doc.add_paragraph()
            run = p.add_run(child.get_text())
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = BLACK
        elif name == "p":
            ctx = "blue" if "blue" in classes else "red" if "red" in classes else "black"
            render_paragraph(doc, child, ctx)
        elif name == "div":
            if "blue" in classes:
                render_callout_block(doc, child, "blue")
                doc.add_paragraph()
            elif "red" in classes:
                render_callout_block(doc, child, "red")
                doc.add_paragraph()
            elif "green" in classes:
                render_callout_block(doc, child, "green")
                doc.add_paragraph()
            else:
                for sub in child.children:
                    if isinstance(sub, Tag):
                        if sub.name == "p":
                            sub_classes = sub.get("class", [])
                            sub_ctx = "blue" if "blue" in sub_classes else "red" if "red" in sub_classes else "green" if "green" in sub_classes else "black"
                            render_paragraph(doc, sub, sub_ctx)
        elif name == "table":
            ctx = "blue" if "blue" in classes else "red" if "red" in classes else "black"
            render_table(doc, child, ctx)
        elif name == "ul":
            render_list(doc, child, "black", ordered=False)
        elif name == "ol":
            render_list(doc, child, "black", ordered=True)
        elif name == "hr":
            doc.add_paragraph("─" * 60)

    doc.save(str(out_path))


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: html_to_colored_docx.py <input.html> [output.docx]")
        sys.exit(1)
    in_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2]) if len(sys.argv) > 2 else in_path.with_suffix(".docx")
    convert(in_path, out_path)
    print(f"Wrote {out_path}")
