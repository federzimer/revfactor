#!/usr/bin/env python3
"""Convert PRESS_RELEASE.md → RevFactor_FIFA_Pitch.docx with readable formatting.
Appends an "About the Source" bio block with Federico's headshot + LinkedIn at the end.
Saves a backup copy to ~/Documents to survive directory sweeps.
"""
import os
import shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(__file__)
SRC = os.path.join(ROOT, "PRESS_RELEASE.md")
OUT = os.path.join(ROOT, "RevFactor_FIFA_Pitch.docx")
BACKUP = "/Users/aaronwhittaker/Claude/_artifacts/RevFactor_FIFA_Pitch.docx"
PHOTO = "/Users/aaronwhittaker/Claude/RevFactor/public/team/federico.jpg"

LINKEDIN = "https://www.linkedin.com/in/federico-zimerman-23ab3345"
WEBSITE = "https://www.revfactor.io"

with open(SRC) as f:
    lines = f.read().splitlines()

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Georgia"
style.font.size = Pt(11)

for raw in lines:
    line = raw.rstrip()
    if not line:
        doc.add_paragraph()
        continue
    if line.startswith("# "):
        p = doc.add_paragraph()
        run = p.add_run(line[2:].strip())
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = "Georgia"
        p.paragraph_format.space_after = Pt(16)
        p.paragraph_format.space_before = Pt(0)
    elif line.startswith("## "):
        p = doc.add_paragraph()
        run = p.add_run(line[3:].strip())
        run.bold = True
        run.font.size = Pt(13)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
    else:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.35
        for run in p.runs:
            run.font.name = "Georgia"
            run.font.size = Pt(11)

# ---- Bio block ----
doc.add_paragraph()
hr = doc.add_paragraph()
hr_run = hr.add_run("___________________________________")
hr_run.font.color.rgb = RGBColor(0xC0, 0xC0, 0xC0)

heading = doc.add_paragraph()
hrun = heading.add_run("About the Source")
hrun.bold = True
hrun.font.size = Pt(13)
heading.paragraph_format.space_before = Pt(12)
heading.paragraph_format.space_after = Pt(8)

if os.path.exists(PHOTO):
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    img_para.add_run().add_picture(PHOTO, width=Inches(1.6))
    img_para.paragraph_format.space_after = Pt(8)

name_p = doc.add_paragraph()
name_run = name_p.add_run("Federico Zimerman")
name_run.bold = True
name_run.font.size = Pt(12)
name_p.paragraph_format.space_after = Pt(2)

title_p = doc.add_paragraph()
title_run = title_p.add_run("Founder, RevFactor")
title_run.italic = True
title_run.font.size = Pt(11)
title_p.paragraph_format.space_after = Pt(8)

bio_text = (
    "Federico Zimerman is the founder of RevFactor, a short-term rental "
    "revenue management service that oversees pricing strategy for 165+ "
    "properties across 24 U.S. states and 56 markets. He works directly "
    "with hosts to apply professional revenue management discipline to STR "
    "portfolios, with a focus on event-driven demand, dynamic pricing, "
    "underpricing diagnostics, and the economics of direct bookings. "
    "Federico is available for interviews and contributed expert commentary."
)
bio_p = doc.add_paragraph(bio_text)
bio_p.paragraph_format.line_spacing = 1.35
bio_p.paragraph_format.space_after = Pt(10)

contact_lines = [
    ("LinkedIn:  ", LINKEDIN),
    ("Website:   ", WEBSITE),
    ("Press:     ", "press@revfactor.io"),
]
for label, value in contact_lines:
    cp = doc.add_paragraph()
    lr = cp.add_run(label)
    lr.bold = True
    lr.font.size = Pt(10)
    vr = cp.add_run(value)
    vr.font.size = Pt(10)
    cp.paragraph_format.space_after = Pt(2)

doc.save(OUT)
# Save a backup outside the synced directory so it survives sweeps
os.makedirs(os.path.dirname(BACKUP), exist_ok=True)
shutil.copy(OUT, BACKUP)
print(OUT)
print(BACKUP)
