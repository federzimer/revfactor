"""
Build a ProCloser-branded docx of the L40 Phase Plan Update.

- ProCloser wordmark logo at top of page 1
- Montserrat headings, Inter body (with Calibri fallback)
- Cyan brand accent (#38BDF8) on H1/H2
- Slate body text (#0F172A)
- Footer: "Prepared by ProCloser.ai  |  RevFactor  |  2026-05-28  |  Page X"

Source: L40_Phase_Plan_Update_2026-05-28.md
Output: L40_Phase_Plan_Update_2026-05-28.docx
"""
from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

import sys

ROOT = Path("/Users/aaronwhittaker/Claude/L40")
# Allow passing a different .md path as argv[1]; fall back to the phase plan
if len(sys.argv) > 1:
    MD = Path(sys.argv[1])
else:
    MD = ROOT / "L40_Phase_Plan_Update_2026-05-28.md"
OUT = MD.with_suffix(".docx")
LOGO = Path("/Users/aaronwhittaker/Claude/ProCloser.ai Website/brand_assests/ProCloser.png")

# Footer date: parse YYYY-MM-DD from the filename (e.g. "L40 - Dashboard - 2026-06-23.md"),
# fall back to today. Keeps the footer aligned with the document instead of a hardcoded date.
import re as _re, datetime as _dt
_m = _re.search(r"(\d{4}-\d{2}-\d{2})", MD.name)
DOC_DATE = _m.group(1) if _m else _dt.date.today().isoformat()

CYAN = RGBColor(0x38, 0xBD, 0xF8)
SLATE_900 = RGBColor(0x0F, 0x17, 0x2A)
SLATE_600 = RGBColor(0x47, 0x55, 0x69)
SLATE_400 = RGBColor(0x94, 0xA3, 0xB8)
SLATE_100 = RGBColor(0xF1, 0xF5, 0xF9)
GREEN = RGBColor(0x10, 0xB9, 0x81)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
RED = RGBColor(0xEF, 0x44, 0x44)

HEADING_FONT = "Montserrat"
BODY_FONT = "Inter"


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="CBD5E1", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tc_pr.append(tcBorders)


def style_paragraph(para, font=BODY_FONT, size=11, color=SLATE_900, bold=False, italic=False):
    for run in para.runs:
        run.font.name = font
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.bold = bold
        run.font.italic = italic
        # Set East Asia font too so Word doesn't auto-substitute
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), font)
        rFonts.set(qn("w:hAnsi"), font)
        rFonts.set(qn("w:cs"), font)


def add_run(para, text, font=BODY_FONT, size=11, color=SLATE_900, bold=False, italic=False):
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    rPr.insert(0, rFonts)
    return run


def add_runs_with_inline_md(para, text, **base_kwargs):
    """Parse inline **bold**, *italic*, `code`, [text](url) → add runs."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            add_run(para, text[pos:m.start()], **base_kwargs)
        token = m.group(0)
        if token.startswith("**"):
            add_run(para, token[2:-2], **{**base_kwargs, "bold": True})
        elif token.startswith("`"):
            add_run(para, token[1:-1], **{**base_kwargs, "font": "Menlo", "color": SLATE_600})
        elif token.startswith("[") and "](" in token:
            label = token[1:token.index("]")]
            add_run(para, label, **{**base_kwargs, "color": CYAN, "bold": True})
        elif token.startswith("*"):
            add_run(para, token[1:-1], **{**base_kwargs, "italic": True})
        pos = m.end()
    if pos < len(text):
        add_run(para, text[pos:], **base_kwargs)


def add_horizontal_rule(doc, color="38BDF8"):
    p = doc.add_paragraph()
    p_pr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    p_pr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)


def add_heading(doc, text, level):
    sizes = {1: 24, 2: 18, 3: 14}
    colors = {1: SLATE_900, 2: CYAN, 3: SLATE_900}
    space_before = {1: 18, 2: 22, 3: 14}
    space_after = {1: 8, 2: 8, 3: 4}

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before[level])
    p.paragraph_format.space_after = Pt(space_after[level])
    p.paragraph_format.keep_with_next = True
    add_run(p, text, font=HEADING_FONT, size=sizes[level], color=colors[level], bold=True)
    if level == 2:
        # underline accent for H2
        p_pr = p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "38BDF8")
        pBdr.append(bottom)
        p_pr.append(pBdr)


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    add_runs_with_inline_md(p, text, font=BODY_FONT, size=11, color=SLATE_900)
    return p


def add_bullet(doc, text, indent_level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * indent_level)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    add_runs_with_inline_md(p, text, font=BODY_FONT, size=11, color=SLATE_900)


def add_ordered_item(doc, text, number):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.35)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    add_run(p, f"{number}.  ", font=BODY_FONT, size=11, color=CYAN, bold=True)
    add_runs_with_inline_md(p, text, font=BODY_FONT, size=11, color=SLATE_900)


def add_md_table(doc, header_cells, body_rows):
    table = doc.add_table(rows=1 + len(body_rows), cols=len(header_cells))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    # Force layout to auto so columns size based on content
    tbl_pr = table._element.find(qn("w:tblPr"))
    if tbl_pr is not None:
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "autofit")
        tbl_pr.append(layout)

    # Header
    for i, h in enumerate(header_cells):
        cell = table.rows[0].cells[i]
        cell.text = ""
        set_cell_shading(cell, "0F172A")
        set_cell_borders(cell, color="0F172A", size="6")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        add_runs_with_inline_md(p, h.strip(), font=HEADING_FONT, size=9,
                                color=RGBColor(0xF8, 0xFA, 0xFC), bold=True)
    # Keep header row with next (so it doesn't orphan at bottom of page)
    tr_pr = table.rows[0]._element.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)
    tbl_header = OxmlElement("w:tblHeader")
    tr_pr.append(tbl_header)

    # Body
    for r_idx, row in enumerate(body_rows):
        zebra = "FFFFFF" if r_idx % 2 == 0 else "F8FAFC"
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            set_cell_shading(cell, zebra)
            set_cell_borders(cell, color="E2E8F0", size="4")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.2
            add_runs_with_inline_md(p, val.strip(), font=BODY_FONT, size=9, color=SLATE_900)
        # Prevent individual rows from splitting across pages
        tr_pr = table.rows[r_idx + 1]._element.get_or_add_trPr()
        cs = OxmlElement("w:cantSplit")
        tr_pr.append(cs)


def parse_alignments(divider_row):
    """Pandoc table align row e.g. |---|---:|--:|"""
    aligns = []
    for cell in divider_row.split("|")[1:-1]:
        s = cell.strip()
        if s.startswith(":") and s.endswith(":"):
            aligns.append("center")
        elif s.endswith(":"):
            aligns.append("right")
        else:
            aligns.append("left")
    return aligns


def split_md_table_row(row):
    """Split a markdown table row, respecting backtick spans and escaped pipes."""
    s = row.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    cells = []
    buf = []
    in_code = False
    j = 0
    while j < len(s):
        c = s[j]
        if c == "\\" and j + 1 < len(s) and s[j + 1] == "|":
            buf.append("|")
            j += 2
            continue
        if c == "`":
            in_code = not in_code
            buf.append(c)
            j += 1
            continue
        if c == "|" and not in_code:
            cells.append("".join(buf).strip())
            buf = []
            j += 1
            continue
        buf.append(c)
        j += 1
    cells.append("".join(buf).strip())
    return cells


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Prepared by ", font=BODY_FONT, size=9, color=SLATE_400)
    add_run(p, "ProCloser.ai", font=HEADING_FONT, size=9, color=CYAN, bold=True)
    add_run(p, f"   ·   RevFactor   ·   {DOC_DATE}   ·   Page ",
            font=BODY_FONT, size=9, color=SLATE_400)

    # page number field
    run = p.add_run()
    run.font.name = BODY_FONT
    run.font.size = Pt(9)
    run.font.color.rgb = SLATE_400
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)


# ----------------------------------------------------------------------
# Build the document
# ----------------------------------------------------------------------

doc = Document()

# Page setup
section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)

# Default style
style = doc.styles["Normal"]
style.font.name = BODY_FONT
style.font.size = Pt(11)
style.font.color.rgb = SLATE_900

# Header — ProCloser logo right-aligned
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hp_run = hp.add_run()
hp_run.add_picture(str(LOGO), width=Inches(1.7))

add_footer(section)

# Parse title + metadata from the markdown header (first H1 + the metadata lines before the first ---)
_md_text = MD.read_text()
_md_lines = _md_text.split("\n")
_doc_title = "L40 Update"
_meta_lines = []
for _i, _ln in enumerate(_md_lines):
    if _ln.startswith("# ") and _doc_title == "L40 Update":
        _doc_title = _ln[2:].strip()
        continue
    if _doc_title != "L40 Update":
        if _ln.strip() == "---":
            break
        if _ln.strip():
            _meta_lines.append(_ln.strip())

# Title block on page 1
title_p = doc.add_paragraph()
title_p.paragraph_format.space_before = Pt(8)
title_p.paragraph_format.space_after = Pt(2)
add_run(title_p, _doc_title, font=HEADING_FONT, size=26, color=SLATE_900, bold=True)

for _line in _meta_lines:
    mp = doc.add_paragraph()
    mp.paragraph_format.space_after = Pt(2)
    add_runs_with_inline_md(mp, _line, font=BODY_FONT, size=11, color=SLATE_600)

add_horizontal_rule(doc)

# ----------------------------------------------------------------------
# Render the markdown content (skip the YAML-style title block we already rendered)
# ----------------------------------------------------------------------

lines = MD.read_text().split("\n")

# Skip past the top metadata block — we re-rendered it nicely above.
# Find the first "---" after the metadata block.
start = 0
for i, line in enumerate(lines):
    # find the second '---' (which closes the metadata) then start rendering after it
    pass
# The doc opens with "# L40 — Phase Plan Update..." then 4 metadata lines, then "---".
# Skip everything up to and including that first "---".
for i, line in enumerate(lines):
    if i > 0 and line.strip() == "---":
        start = i + 1
        break
lines = lines[start:]

i = 0
while i < len(lines):
    line = lines[i].rstrip()

    if not line.strip():
        i += 1
        continue

    if line.strip() == "---":
        add_horizontal_rule(doc, color="CBD5E1")
        i += 1
        continue

    if line.startswith("## "):
        add_heading(doc, line[3:].strip(), 2)
        i += 1
        continue
    if line.startswith("### "):
        add_heading(doc, line[4:].strip(), 3)
        i += 1
        continue
    if line.startswith("# "):
        add_heading(doc, line[2:].strip(), 1)
        i += 1
        continue

    # Tables — pandoc-style: header row | divider | body rows
    if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1]):
        header = split_md_table_row(line)
        body = []
        j = i + 2
        while j < len(lines) and lines[j].startswith("|"):
            body.append(split_md_table_row(lines[j]))
            j += 1
        add_md_table(doc, header, body)
        # small spacer after table
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(2)
        i = j
        continue

    # Numbered lists
    m_num = re.match(r"^(\d+)\.\s+(.*)$", line)
    if m_num:
        add_ordered_item(doc, m_num.group(2), m_num.group(1))
        i += 1
        continue

    # Bullets (- or *)
    if line.startswith("- ") or line.startswith("* "):
        add_bullet(doc, line[2:])
        i += 1
        continue

    # Blockquotes
    if line.startswith("> "):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(6)
        # left border for blockquote feel
        p_pr = p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "12")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), "38BDF8")
        pBdr.append(left)
        p_pr.append(pBdr)
        add_runs_with_inline_md(p, line[2:], font=BODY_FONT, size=11, color=SLATE_600, italic=True)
        i += 1
        continue

    # Plain paragraph
    add_paragraph(doc, line)
    i += 1

doc.save(str(OUT))
print(f"Wrote {OUT}")
