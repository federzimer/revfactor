#!/usr/bin/env python3
"""
Post-process revfactor-content-brain.docx to:
1. Find any paragraph containing [ROUND-2] marker
2. Strip the marker
3. Color the entire paragraph blue
4. Also color all paragraphs from "Round-2 Deep Mine" heading onwards
"""

from docx import Document
from docx.shared import RGBColor

BLUE = RGBColor(0x1A, 0x4F, 0xBF)
GREEN = RGBColor(0x0A, 0x7A, 0x3C)
DOCX = "/Users/aaronwhittaker/Claude/RevFactor/docs/revfactor-content-brain.docx"

doc = Document(DOCX)

in_round2_section = False
in_delta_section = False
blue_count = 0
green_count = 0

for para in doc.paragraphs:
    text = para.text or ""

    # Section transitions
    if "Round-2 Deep Mine" in text:
        in_round2_section = True
        in_delta_section = False  # exiting any delta section if we re-enter round-2
    if "Updates Since v1 Sent to GetCito" in text or "19. Updates Since" in text:
        in_round2_section = False
        in_delta_section = True

    has_round2_marker = "[ROUND-2]" in text or "[ROUND-2 START" in text
    has_delta_marker = "[DELTA]" in text

    target_color = None
    if in_delta_section or has_delta_marker:
        target_color = GREEN
    elif in_round2_section or has_round2_marker:
        target_color = BLUE

    if target_color:
        # Strip marker text from runs and color them
        for run in para.runs:
            if run.text:
                run.text = (run.text
                    .replace("[ROUND-2 START — added 2026-04-30 from a deeper pass through 3 long-form YouTube interviews previously flagged as not yet mined: Craft Stays (\"75 Properties Later: Why a Delisted Airbnb Made Federico\"), Catchup with the Carlyles (\"A Journey through Hospitality\"), and Life of Flow (\"How to Build a Profitable Airbnb Business\"). Color-coded blue to distinguish from the original brain content above.]", "*Round-2 deep-mine content (2026-04-30) — added from Craft Stays, Catchup with the Carlyles, and Life of Flow long-form interviews. Color-coded blue to distinguish from the original brain content above.*")
                    .replace("[ROUND-2] ", "")
                    .replace("[ROUND-2]", "")
                    .replace("[DELTA] ", "")
                    .replace("[DELTA]", "")
                )
                run.font.color.rgb = target_color
        if target_color == GREEN:
            green_count += 1
        else:
            blue_count += 1
    else:
        colored_count_unset = 0  # original content stays default

# For backward compat in print
colored_count = blue_count + green_count

# Also walk tables in case any are inside Round-2 section (they shouldn't be, but safety)
# Tables before the round-2 section should remain default-colored
section_18_started = False
for tbl in doc.tables:
    # Check if this table is after section 18 (best-effort: tables after the heading)
    # python-docx table iteration order matches document order, but we don't have a
    # reliable way to map table → section. For now, leave tables uncolored — the
    # round-2 content doesn't include critical tables.
    pass

doc.save(DOCX)
print(f"Colored {blue_count} paragraphs blue + {green_count} paragraphs green. Saved to {DOCX}")
