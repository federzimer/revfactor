from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# ── Style helpers ─────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x13, 0x34, 0x2D)   # cedar
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x3F, 0x26, 0x1F)   # tobacco
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x5D, 0x6D, 0x59)   # moss
    return p

def body(text, italic=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.italic = italic
    run.font.color.rgb = RGBColor(0x3F, 0x26, 0x1F)
    return p

def prefilled(label, text):
    """Green tick label + body text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label + " ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor(0x5D, 0x6D, 0x59)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor(0x3F, 0x26, 0x1F)

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x3F, 0x26, 0x1F)

def numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x3F, 0x26, 0x1F)

def field_label(label, answer=""):
    """Bold label then answer on same para."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label + " ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor(0x76, 0x57, 0x4C)  # walnut
    if answer:
        r2 = p.add_run(answer)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = RGBColor(0x3F, 0x26, 0x1F)

def additions_header():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("Corrections / Additions:")
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x13, 0x34, 0x2D)

def divider():
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("CLIENT ONBOARDING QUESTIONNAIRE")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x13, 0x34, 0x2D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("SEO + GEO Combined Optimization")
rs.font.size = Pt(12)
rs.font.color.rgb = RGBColor(0x5D, 0x6D, 0x59)

for label, val in [("Prepared for:", "RevFactor"), ("Prepared by:", "GetCito Pvt. Ltd."), ("Completed:", "April 2026")]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + " ")
    r1.bold = True; r1.font.size = Pt(10.5)
    r2 = p.add_run(val)
    r2.font.size = Pt(10.5)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Brand, Product & Positioning")

# 1.1
h2("1.1. Company & Founding Story")
prefilled("[✔ Pre-filled]:",
    " RevFactor is a specialized revenue management and pricing agency for short‑term rental (STR) "
    "hosts and property managers, led by Federico Zimerman. Unlike generic pricing software, RevFactor "
    "layers AI‑assisted data analysis with expert human strategy to maximize rental income and elevate "
    "what it calls \"Inspired Stays.\" As of March 2025, the company manages 100+ properties across "
    "diverse U.S. markets.\n\n"
    "Positioning itself as the \"intelligent assistant\" for property owners, RevFactor focuses on revenue "
    "intelligence, not just tools. Its approach combines dynamic pricing platforms like PriceLabs with "
    "proprietary frameworks such as the Market Position Score, Pricing Health Index, and real‑time "
    "competitive analysis. The result: volatility is transformed into growth, and hosts gain a clear, "
    "strategic edge in competitive rental markets.")

additions_header()
body(
    "RevFactor currently operates as Blackberry Hospitality (launched August 2023 after Federico parted "
    "ways with his first management company on good terms). He previously scaled a prior management company "
    "to 75 properties before relaunching with a quality-first focus, now managing across 16 U.S. states.\n\n"
    "Federico's authentic mission statement, in his own words: \"Provide peace of mind to owners and a "
    "five-star stay to guests.\" His positioning is grounded in his 10-year airline revenue management "
    "background — a genuinely unique credential in the STR space that should be featured prominently.\n\n"
    "The Interest + Reliability + Positioning framework (derived from airline yield management) is "
    "Federico's actual three-pillar methodology and should be incorporated into all content describing "
    "how RevFactor works."
)

# 1.2
h2("1.2. Core Service Offerings")
prefilled("[✔ Pre-filled]:", " RevFactor provides a managed pricing service. Key pillars include:")
for item in [
    "Revenue management for short‑term rentals (Airbnb, Vrbo, etc.)",
    "Dynamic pricing strategy (demand‑responsive, event‑aware, market‑optimized)",
    "Competitive set analysis & market positioning",
    "Pricing health audit & revenue opportunity assessment",
    "Daily monitoring & strategic adjustments",
    "Monthly reporting & performance tracking",
    "Flat monthly fee (designed to pay for itself)",
]:
    bullet(item)

body("Additional offerings to highlight:", indent=False)
for item in [
    "Multi-channel OTA management: Airbnb, VRBO, Booking.com, Google Vacation Rentals, and Marriott Homes & Villas by Bonvoy (RevFactor is an approved Marriott partner)",
    "New listing launch strategy — market penetration pricing to accelerate review accumulation and Airbnb ranking",
    "Documentation and dispute support systems (checklists, photo proof, Airbnb claim defense)",
    "Direct booking strategy and channel diversification — Federico's stated goal is reducing Airbnb dependency from 85% to 50% of bookings for clients",
    "Pricing: flat fee of $320/month per property, with volume discounts for multiple properties",
]:
    bullet(item)

field_label("Are there specific offerings that we have missed, or should we highlight?",
            " _________________________________________________")

# 1.3
h2("1.3. Service Description")
prefilled("[✔ Pre-filled]:", " RevFactor provides a managed pricing service. Key pillars include:")
for label, desc in [
    ("Dynamic Pricing:", "Real-time rate adjustments based on local demand, events, and seasonality."),
    ("Market Analysis:", "Benchmarking property performance against local competitors."),
    ("Occupancy Optimization:", "Balancing nightly rates with high booking volume to ensure maximum yield."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label + " ")
    r1.bold = True; r1.font.size = Pt(10.5)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5)

body("\nAdditional proprietary elements to highlight:")
body(
    "Federico's differentiating methodology is the Interest + Reliability + Positioning framework, "
    "derived from 10 years in airline tariff and revenue management:\n"
    "• Interest: Conversion rate optimization — listing photos (strategic ordering), descriptions, "
    "amenity completeness, Airbnb professional hosting tools\n"
    "• Reliability: Review management and guest experience — documentation systems, proactive "
    "communication, recovery protocols\n"
    "• Positioning: Daily pricing relative to competitive set, using pacing data and RevPAR "
    "(not just ADR) as the primary success metric\n\n"
    "RevFactor's approach: understand the math first, then use tools to automate it. "
    "\"The important thing is to understand the math\" — tools like PriceLabs execute the strategy "
    "Federico designs."
)
field_label('Are there specific proprietary tools or "RevFactor-exclusive" algorithms we should highlight?',
            " _________________________________________________")

# 1.4
h2("1.4. Core Value Proposition (USP)")
prefilled("[✔ Pre-filled]:",
    " RevFactor's core value proposition is simple yet powerful: \"Intelligent Pricing for Inspired "
    "Stays.\" Unlike standard automated tools such as PriceLabs or Airbnb's Smart Pricing, RevFactor "
    "adds a strategic human layer that consistently captures margins software alone misses, often 15–20% "
    "during peak events or shoulder seasons.\n\n"
    "Its most defensible differentiator lies in this hybrid model: the only service that combines "
    "PriceLabs' dynamic pricing engine with daily expert review and proprietary revenue intelligence "
    "frameworks. With tools like the Market Position Score, Pricing Health Index, and real‑time "
    "competitive analysis, RevFactor delivers a proven revenue lift (+18% vs. comp set) while operating "
    "on a transparent flat‑fee model that never takes a percentage of client revenue.")

additions_header()
body("Additional USPs to highlight:")
for item in [
    "10-year airline revenue management background — Federico trained in yield management, customer segmentation, and tariff strategy at an international airline. No other STR revenue management service has this pedigree.",
    "The Two-Client Philosophy — \"We're handling the two most expensive assets people have — the owner's property and the guest's time. We don't want to mess that up.\"",
    "Flat fee ($320/month per property), not percentage of revenue — hosts keep 100% of what they earn above the fee.",
    "Operator credibility — Federico actively manages 75+ properties. Results are proven at scale, not theoretical.",
    "Marriott Homes & Villas partnership — access to higher-income, higher-trust travelers not reachable through Airbnb/VRBO alone.",
    "Documentation and dispute protection — proven case study: $900+ guest refund claim prevented through checklist and photo documentation systems.",
]:
    bullet(item)
field_label("Please list any additional USPs you wish us to highlight:",
            " _________________________________________________")

# 1.5
h2("1.5. Are there any claims, associations, or customer types RevFactor must NEVER be linked with?")
for item in [
    "Full property management services (cleaning, maintenance, guest check-in, operations) — RevFactor is revenue management only",
    "Passive income / \"set and forget\" framing — Federico explicitly rejects this in every interview",
    "Guaranteed revenue promises — results vary by market, property, and season",
    "\"Get rich with Airbnb\" content or influencer-style STR courses",
    "Boutique hotels or commercial hospitality — RevFactor serves STR hosts specifically",
    "Latin American market operations — RevFactor is U.S.-only (16 states currently)",
]:
    bullet(item)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Target Audience & ICP (Ideal Customer Profile)")

h2("2.1. Primary Personas")
prefilled("[✔ Pre-filled]:", "")
for persona in [
    ("Persona 1: The Multi-Unit Property Manager (PMC).",
     " Managing 10–100+ units, looking for professionalized revenue management to scale without increasing headcount."),
    ("Persona 2: The Boutique Hotel Owner.",
     " Independent operators who lack a full-time revenue manager but need to compete with corporate chains."),
    ("Persona 3: High-End STR Investors.",
     " Owners of luxury properties where a single missed booking represents a significant loss."),
    ("Persona 4: The Solo STR Host.",
     " 1–3 properties, side hustle or full‑time, wants to maximize income without spending hours on pricing."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(persona[0])
    r1.bold = True; r1.font.size = Pt(10.5)
    r2 = p.add_run(persona[1])
    r2.font.size = Pt(10.5)

additions_header()
body(
    "Based on Federico's stated positioning (2025–2026 interviews), his ICP has evolved after his Phase 2 "
    "business pivot (Blackberry Hospitality, August 2023). He explicitly works with quality-first, "
    "hospitality-focused hosts who \"give him the freedom to make sure everything is perfect.\"\n\n"
    "Strongest ICP segment: the Growing Quality Operator — managing 3–15 properties, has tried DIY "
    "pricing tools, wants expert management without giving up a revenue percentage to a full property manager.\n\n"
    "High-priority trigger audience: hosts who just scaled to their 3rd–5th property and the DIY pricing "
    "approach is breaking down across multiple listings and markets.\n\n"
    "The Boutique Hotel persona should be treated as secondary — Federico's operational experience and "
    "content are 100% STR-focused."
)

h2("2.2. Geographic Focus")
prefilled("[✔ Pre-filled]:", " Global capability with a current strong focus on U.S. and Latin American markets.")
additions_header()
body("Confirmed active U.S. markets (from interviews): West Michigan, Upstate New York, Fort Worth/Texas, "
     "Texas Hill Country, Scottsdale AZ, Smoky Mountains, Tahoe, Poconos.\n\n"
     "Federico's data-edge markets (deepest operational history): Fort Worth TX, West Michigan, Upstate NY.\n\n"
     "Note: Latin American market operations are not confirmed. RevFactor currently operates across 16 U.S. states.")
field_label("Are there specific high-demand markets where RevFactor has a 'data edge'?",
            " West Michigan, Tahoe, Smoky Mountains, Upstate New York, Washington (pre-filled) — "
            "also confirmed: Fort Worth TX, Scottsdale AZ, Texas Hill Country")

h2("2.3. Top 3 Problems / Frustrations")
prefilled("[✔ Pre-filled examples from website]:", "")
for item in [
    "\"Pricing without strategy is just guessing\"",
    "\"Leaving money on the table during peak demand periods\"",
    "\"Gap nights and inconsistent occupancy eating into profits\"",
    "\"Reacting to competitors blindly\"",
    "\"Missing local events, seasonal shifts, and market changes\"",
]:
    bullet(item)

body("\nTop 3 based on Federico's actual client conversations:")
numbered("ADR obsession over RevPAR — Hosts fixate on nightly rate without tracking actual revenue generated. Federico: \"I asked her how much she was selling. Silence.\" They're optimizing the wrong metric.")
numbered("Set-and-forget tool failure at scale — Pricing tools work passably for 1–2 properties. At 5+, the tool runs but nobody's watching pacing, conversion rates, or competitive shifts. Revenue quietly leaks.")
numbered("Missing the new listing window — Most hosts launch at market rate and wait. Federico's market penetration strategy builds reviews fast and signals quality to the algorithm early. Most hosts never recover from a poor launch.")

h2("2.4. Purchase Triggers & Objections")
body("Triggers:", italic=False)
numbered("A specific, visible revenue miss — a local event passed and they didn't adjust. A similar property nearby made 3x what they did. The pain is concrete.")
numbered("Just added their 3rd–5th property — DIY pricing that worked for 1–2 is now unmanageable. Too many variables, not enough time.")
numbered("Competitor is outperforming them with no obvious reason — Federico's audit reveals photo ordering, amenity gaps, pricing cadence problems — all fixable, none obvious without expert eyes.")
body("\nObjections:")
numbered("\"I already pay for PriceLabs / Wheelhouse — isn't that enough?\" → Tools are the engine, not the strategy. RevFactor is the driver.")
numbered("\"I don't want to pay a percentage of my revenue\" → RevFactor charges a flat fee ($320/month per property), not a cut of bookings.")
numbered("\"I need to stay in control of my pricing\" → Federico provides strategy and daily execution while keeping owners informed. Peace of mind, not loss of control.")

h2("2.5. Typical Property Types")
body("Entire homes (primary), cabins and rural retreats (Smoky Mountains, Texas Hill Country, Upstate NY), "
     "design-forward/boutique STRs, urban condos/apartments, properties meeting Marriott Homes & Villas quality criteria.\n\n"
     "Federico's Phase 2 focus: quality-first hosts who invest in design, amenities, and guest experience.")

h2("2.6. Average Nightly Rate")
body("$150–$500/night is the core range based on Federico's portfolio discussions. Luxury outliers exist "
     "(Marriott Homes & Villas program properties command premium rates). Confirm the actual client range "
     "with Federico for marketing purposes.")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Competitive Landscape")

h2("3.1. Primary Competitors")
prefilled("[✔ Pre-filled]:", " SaaS-only tools like PriceLabs, Beyond (formerly Beyond Pricing), and Wheelhouse.")
prefilled("[✔ Pre-filled]:", " Full-service management firms that bundle pricing with cleaning/operations (e.g., Vacasa).")
field_label("What makes a client switch from a tool like PriceLabs to a service like RevFactor?")
body(
    "Federico's answer: Clients switch when a specific revenue miss or bad season makes the \"set and forget\" "
    "problem visible. The tool runs but nobody's watching — no pacing analysis, no conversion rate tracking, "
    "no competitive adjustment. When they realize the tool is doing something but they can't explain why "
    "and results are inconsistent, they want a human who can. The switch accelerates when managing 5+ "
    "properties and the cognitive load becomes untenable."
)

h2("3.2. What RevFactor Does That None of These Competitors Can Claim")
for item in [
    "10-year airline revenue management background — yield management, customer segmentation, and tariff strategy applied to STRs. Genuinely unique in this market.",
    "Flat fee ($320/month per property), not percentage of revenue — full property managers take 20–40% of revenue. Hosts keep 100% of what they earn.",
    "Active operator at scale — Federico manages 75+ properties. Not a consultant. Proven at scale.",
    "Multi-channel OTA management — Airbnb, VRBO, Booking.com, Google Vacation Rentals, Marriott Homes & Villas. Pricing tools only optimize one channel.",
    "Documentation and dispute protection — proven: $900+ refund claim prevented through checklist and photo proof systems.",
    "Marriott Homes & Villas partner — access to Marriott's high-income traveler base.",
]:
    bullet(item)

h2("3.3. Competitor Content / Case Studies to Reference")
body("Suggested: PriceLabs blog (pricing education), Rental Scale-Up (STR industry editorial — AI cites this "
     "heavily), Beyond's Revenue Manager's Guide series.")
field_label("Share links if available:", " _________________________________________________")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Generative Engine Optimization (GEO) & Search")

h2("4.1. High-Value AI Prompts")
body("We want AI engines (ChatGPT, Perplexity, Gemini) to recommend RevFactor when users ask:")
for item in [
    "\"How can I increase my Airbnb revenue by 20% this season?\"",
    "\"What are the best dynamic pricing services for short-term rentals?\"",
    "\"Who is the top-rated revenue manager for boutique hotels?\"",
]:
    bullet(item)

body("\nAdditional high-priority AI prompts (based on Peec.ai live data — 75 tracked prompts):")
for item in [
    "\"Who offers full-service revenue management for Airbnb hosts?\"",
    "\"What companies offer done-for-you Airbnb pricing services?\"",
    "\"What are the best revenue management companies for short-term rentals?\"",
    "\"Who provides pricing strategy consulting for short-term rentals?\"",
    "\"How do I outsource pricing for my short-term rental?\"",
    "\"What is the difference between a pricing tool and a revenue management service?\"",
    "\"Best alternatives to PriceLabs for Airbnb hosts\"",
    "\"How to increase STR revenue without raising nightly rates\"",
]:
    bullet(item)

body("\nIndustry pain points for AI discovery:")
for item in [
    "Airbnb algorithm changes and ranking factors",
    "Pacing and booking window strategy",
    "Low season / shoulder season occupancy strategies",
    "New listing launch and market penetration strategy",
    "Multi-OTA channel management",
    "Airbnb dispute resolution and documentation systems",
    "ADR vs RevPAR — understanding the right metrics",
]:
    bullet(item)

h2("4.2. Credibility & Social Proof")
prefilled("[✔ Pre-filled]:",
    " Feature appearances on industry podcasts (e.g., No Vacancy with Natalie Palmer) and speaking "
    "engagements at summits like Level Up Your Listing.")

additions_header()
body("Verified podcast appearances (from transcripts):")
for item in [
    "Patryk Real Estate Show — \"Advanced Airbnb Pricing and Marketing Strategies\" (confirmed)",
    "Live Let Thrive with Steven Suarez (confirmed — Federico is a regular collaborator)",
    "Life of Flow podcast (confirmed)",
    "Catchup with the Carlyles — Sarah and Emily's podcast; Federico manages their properties (confirmed)",
    "For Your Own Good podcast (confirmed)",
    "Craft Stays podcast (confirmed)",
]:
    bullet(item)

body("\nVerified speaking appearance:")
bullet("Airbnb's first-ever Property Manager Summit — Airbnb HQ. 200 property managers invited. "
       "Federico confirmed as one of the invited speakers. Quote: \"In two days I'm flying to Airbnb "
       "headquarters. They're doing their first summit with property managers ever. 200 invited. I'm "
       "going there and from those 200, they invited me to be a speaker.\" Significant credibility milestone.")

body("\nRevenue lift case studies to develop:")
numbered("The $900 Refund Defense — Guest demanded a $900+ refund for alleged damage. RevFactor's photo documentation and timestamped checklist system disputed the claim. Refund denied.")
numbered("The Delisted Property Recovery — Client property removed from Airbnb. RevFactor activated Booking.com, Google Vacation Rentals, and direct bookings while challenging. Property recovered. Revenue maintained during blackout period.")
numbered("40% Occupancy, 3x Revenue — Properties that run 40% occupancy in slow season but generate 3x more revenue than comparable listings. Specific property data needed for full case study.")
field_label("Can you provide 2–3 specific Revenue Lift case studies?", " _________________________________________________")

divider()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Technical & Integration")

h2("5.1. Tech Stack Compatibility")
field_label("Which Property Management Systems (PMS) do you work with most frequently?")
body("PMS / tools confirmed from interviews: PriceLabs (primary dynamic pricing tool), "
     "Hospitable (confirmed referenced), channel managers for multi-OTA management.\n\n"
     "Verify directly with Federico: Guesty, Hostaway, OwnerRez, Lodgify — likely compatible "
     "but not confirmed in transcripts.")
field_label("Do you require API access to the client's listings, or do you manage pricing via third-party software?")
body("RevFactor manages pricing via PriceLabs (or similar tool) with client-provided account access. "
     "Federico's team does daily expert review and adjustments on top of the tool's automated recommendations.")

h2("5.2. Critical Business Objectives (Next 6 Months)")
numbered("Reduce Airbnb dependency from 85% → 50% of bookings — building direct booking channels, Booking.com, Google Vacation Rentals, and Marriott Homes & Villas. Core operational and risk-management priority.")
numbered("Establish AI search visibility — currently 0% visibility across all tracked prompts on ChatGPT, Perplexity, Gemini, Google AI Overview, and Google AI Mode (confirmed via Peec.ai, April 2026). Goal: appear for \"done-for-you STR revenue management\" queries within 6 months.")
field_label("Objective 3:", " _________________________________________________")

h2("5.3. AI Visibility Gaps")
prefilled("[✔ Pre-filled]:",
    " Many service‑based STR businesses have low visibility in AI engines (ChatGPT, Perplexity, Gemini). "
    "When someone asks: \"Best revenue management service for Airbnb hosts\" or \"How to increase STR "
    "revenue without raising prices\" – RevFactor likely does not appear.")
field_label("Have you ever searched for your own services in ChatGPT or Perplexity? What did you find?",
            " We don't show up at all.")

additions_header()
body(
    "Confirmed via live Peec.ai data (April 2026): RevFactor has 0% AI visibility across all 75 tracked "
    "prompts on all 5 AI engines. This is not estimated — it is confirmed with live tracking data.\n\n"
    "Competitor visibility: PriceLabs 86–98% | Wheelhouse 59–94% | RevFactor 0% across all engines.\n\n"
    "Primary root cause: RevFactor's website is a client-side React SPA that renders blank HTML to crawlers. "
    "AI engines cannot index its content. No sitemap, schema markup, or editorial presence on the domains "
    "AI citation sources pull from.\n\n"
    "Citation gap domains — AI pulls these but RevFactor never appears:\n"
    "• pricelabs.co (51% retrieved, citation rate 1.89)\n"
    "• hostaway.com (41%, citation rate 1.58)\n"
    "• reddit.com (22%, citation rate 1.15)\n"
    "• hoteltechreport.com (18%, citation rate 1.29)\n"
    "• rentalscaleup.com (8%, citation rate 1.05)"
)

h2("5.4. Brand Voice")
body("Filled from persona guide (126 Instagram transcripts + 7 YouTube interviews, ~53,000 words of Federico's own speech):\n")
for label, val in [
    ("Tone:", "Calm, analytical, expert without arrogance. Genuinely enthusiastic about hospitality — not hype-y."),
    ("Structure:", "Short punchy opener → medium explanatory build → longer actionable close. Natural list structure: \"Number one... and the other way is...\""),
    ("Signature phrases (use):", "\"Let me show you [x]\" / \"Here's the thing...\" / \"The important thing is to understand [principle]\" / \"Keep in mind that...\" / \"So, again, [restatement]\""),
    ("Never use:", "\"Leverage your assets\" / \"Synergistic\" / \"Circle back\" / \"Revenue optimization\" / Any passive income framing"),
    ("Teaching arc:", "Problem → Principle → Specific example → Actionable step → Principle restatement"),
    ("Core belief:", "\"This is a people business, not just real estate. We're handling the two most expensive assets people have — their property and their time.\""),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label + " ")
    r1.bold = True; r1.font.size = Pt(10.5)
    r2 = p.add_run(val)
    r2.font.size = Pt(10.5)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# SUBMISSION
# ══════════════════════════════════════════════════════════════════════════════
h1("Submission")
body(
    "Thank you for providing these insights. This information will be used to calibrate our Generative "
    "Engine Optimization (GEO) strategy to ensure RevFactor is identified as the \"Authority\" in "
    "hospitality revenue management across all AI platforms."
)
doc.add_paragraph()
for line in ["Confirmed by:", "Name: _____________________", "Role: _____________________", "Date: _____________________"]:
    p = doc.add_paragraph(line)
    p.paragraph_format.space_after = Pt(6)
    p.runs[0].font.size = Pt(10.5)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/Users/aaronwhittaker/Claude/RevFactor/Files/RevFactor_Onboarding_Questionnaire_COMPLETED.docx"
doc.save(out)
print(f"Saved: {out}")
